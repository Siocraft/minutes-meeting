import os
import sys
from openai import OpenAI
from docx import Document
from pydub import AudioSegment
from tqdm import tqdm

if not os.environ.get("OPENAI_API_KEY"):
  print("ERROR: Set the OPENAI_API_KEY environment variable.")
  print("export OPENAI_API_KEY='your-api-key'")
  exit()

client = OpenAI(
  api_key=os.environ.get("OPENAI_API_KEY"),
)

KAIROS_PREFIX = "\033[34m[Kairos]\033[0m"

def show_loading():
  for _ in tqdm(range(100), desc="\033[54m[Kairos]\033[0m"):
    pass

def divide_audio(input_file, segment_duration=60):
  audio = AudioSegment.from_wav(input_file)
  print(f"{KAIROS_PREFIX} Dividing audio file into chunks of {segment_duration} seconds...")
  segment_length_ms = segment_duration * 1000

  for i, start_time in enumerate(range(0, len(audio), segment_length_ms)):
    segment = audio[start_time:start_time + segment_length_ms]
    output_file = f"segments/segment_{i + 1}.wav"
    segment.export(output_file, format="wav")

  global segments_quantity
  segments_quantity = i + 1

  print(f"{KAIROS_PREFIX} Audio file divided into \033[35m{segments_quantity} chunks\033[0m.")

def transcribe_single_audio(audio_file_path):
  show_loading()
  with open(audio_file_path, 'rb') as audio_file:
    transcription = client.audio.transcriptions.create(
      model="whisper-1",
      file=audio_file,
      prompt="TCS, FY (Financial Year), SAP",
      language="en"
    )
  print(f"{KAIROS_PREFIX} Audio chunk file transcribed \033[35msuccessfully\033[0m.")
  return transcription.text

def transcribe_audio(audio_file_path, audio_text_file_path="audio_text.txt"):
  divide_audio(audio_file_path)
  audio_chunks = [open(f"segments/segment_{i + 1}.wav", "rb").read() for i in range(segments_quantity)]
  audio_text_file = open(audio_text_file_path, "w")
  transcriptions = []
  for i, chunk in enumerate(audio_chunks):
    print(f"{KAIROS_PREFIX} Transcribing audio chunk {i + 1}")
    with open(f"segments/segment_{i+1}.wav", 'wb') as chunk_file:
      chunk_file.write(chunk)
    transcription = transcribe_single_audio(f"segments/segment_{i+1}.wav")
    transcriptions.append(transcription)
  
  all_transcriptions = ' '.join(transcriptions)
  audio_text_file.write(all_transcriptions)
  print(f"\n\n{KAIROS_PREFIX} Complete audio transcribed \033[35msuccessfully\033[0m.\n\n")
  return all_transcriptions

def abstract_summary_extraction(transcription):
  print(f"{KAIROS_PREFIX} Extracting abstract summary...")
  show_loading()
  response = client.chat.completions.create(
    model="gpt-4",
    temperature=0,
    messages=[
      {
        "role": "system",
        "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
      },
      {
        "role": "user",
        "content": transcription
      }
    ]
  )
  print(f"{KAIROS_PREFIX} Abstract summary extracted \033[35msuccessfully\033[0m.")
  return response.choices[0].message.content

def key_points_extraction(transcription):
  print(f"{KAIROS_PREFIX} Extracting key points...")
  show_loading()
  response = client.chat.completions.create(
    model="gpt-4",
    temperature=0,
    messages=[
      {
        "role": "system",
        "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
      },
      {
        "role": "user",
        "content": transcription
      }
    ]
  )
  print(f"{KAIROS_PREFIX} Key points extracted \033[35msuccessfully\033[0m.")
  return response.choices[0].message.content

def action_item_extraction(transcription):
  print(f"{KAIROS_PREFIX} Extracting action items...")
  show_loading()
  response = client.chat.completions.create(
    model="gpt-4",
    temperature=0,
    messages=[
      {
        "role": "system",
        "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
      },
      {
        "role": "user",
        "content": transcription
      }
    ]
  )
  print(f"{KAIROS_PREFIX} Action items extracted \033[35msuccessfully\033[0m.")
  return response.choices[0].message.content

def sentiment_analysis(transcription):
  print(f"{KAIROS_PREFIX} Analyzing sentiment...")
  show_loading()
  response = client.chat.completions.create(
    model="gpt-4",
    temperature=0,
    messages=[
      {
        "role": "system",
        "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
      },
      {
        "role": "user",
        "content": transcription
      }
    ]
  )
  print(f"{KAIROS_PREFIX} Sentiment analyzed \033[35msuccessfully\033[0m.")
  return response.choices[0].message.content

def meeting_minutes(transcription):
  abstract_summary = abstract_summary_extraction(transcription)
  key_points = key_points_extraction(transcription)
  # action_items = action_item_extraction(transcription)
  # sentiment = sentiment_analysis(transcription)
  return {
    'abstract_summary': abstract_summary,
    'key_points': key_points,
    # 'action_items': action_items,
    # 'sentiment': sentiment
  }

def text_to_speech(text, filename):
  print(f"{KAIROS_PREFIX} Converting text to speech...")
  show_loading()
  response = client.audio.speech.create(
    model="tts-1",
    voice="alloy",
    input=text
  )

  response.stream_to_file(filename)

  print(f"{KAIROS_PREFIX} Text converted to speech \033[35msuccessfully\033[0m.")

def save_as_docx(minutes, filename):
  print(f"\n\n{KAIROS_PREFIX} Saving meeting minutes as a Word document in \033[35m.docx format\033[0m")
  doc = Document()
  for key, value in minutes.items():
    heading = ' '.join(word.capitalize() for word in key.split('_'))
    doc.add_heading(heading, level=1)
    doc.add_paragraph(value)
    doc.add_paragraph()
  doc.save(filename)
  print(f"\n\n{KAIROS_PREFIX} Meeting minutes saved \033[35msuccessfully\033[0m.")

def convert_m4a_to_wav(audio_file_path):
  print(f"{KAIROS_PREFIX} Converting audio file to WAV format...")
  os.system(f"ffmpeg -i {audio_file_path} -acodec pcm_s16le -ac 1 -ar 16000 {audio_file_path.replace('.m4a', '.wav')}")
  print(f"{KAIROS_PREFIX} Audio file converted to WAV format \033[35msuccessfully\033[0m.")
  return audio_file_path.replace('.m4a', '.wav')

def get_audio_file_path():
  if '-a' in sys.argv:
    audio_file_path = sys.argv[sys.argv.index('-a') + 1]
    if audio_file_path.endswith('.m4a'):
      audio_file_path = convert_m4a_to_wav(audio_file_path)
    return audio_file_path

  audio_file_path = input(f"{KAIROS_PREFIX} Enter the path of the audio file: ")
  if audio_file_path.endswith('.m4a'):
    audio_file_path = convert_m4a_to_wav(audio_file_path)
  return audio_file_path

def reset_state():
  os.system('rm -rf texts')
  os.system('rm -rf segments')
  os.system('rm -rf speech')
  os.system('mkdir texts')
  os.system('mkdir segments')
  os.system('mkdir speech')
  os.system('clear')

print(f"{KAIROS_PREFIX} Welcome to \033[34mKairos\033[0m - Your AI Meeting Assistant")
print(f"{KAIROS_PREFIX} Kairos helps you taking minutes of any \033[35mmeeting\033[0m.\n\n")

segments_quantity = 0

audio_file_path = get_audio_file_path()
audio_text_file_path = audio_file_path.replace('.wav', '.txt').replace('audios', 'texts')

speech_file = audio_text_file_path.replace('.txt', '.mp3').replace('texts', 'speech')

transcription = transcribe_audio(audio_file_path, audio_text_file_path)
minutes_of_the_meeting = meeting_minutes(transcription)

text_to_speech(minutes_of_the_meeting['abstract_summary'], speech_file)

docx_filename = audio_text_file_path.replace('.txt', '.docx')
save_as_docx(minutes_of_the_meeting, docx_filename)

print(f"\n\n{KAIROS_PREFIX} Time to generate the video.\n\n")

print("\033[34m=========================== BEGIN OF ABSTRACT ===========================\033[0m")
print("\033[34m" + minutes_of_the_meeting['abstract_summary'] + "\033[34m")
print("\033[34m============================ END OF ABSTRACT ============================\033[0m")

print("\n\n\033[34m=========================== BEGIN OF KEY POINTS ===========================\033[0m")
print("\033[34m" + minutes_of_the_meeting['key_points'] + "\033[34m")
print("\033[34m============================ END OF KEY POINTS ============================\033[0m")